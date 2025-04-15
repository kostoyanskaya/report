from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import pandas as pd

from core.constants import (
    CLASSID_FEDERAL,
    CLASSID_REGIONAL,
    FEDERAL_PRESENT_TEXT,
    INTRO_TEXT,
    NO_FEDERAL_TEXT,
    SUMMARY_TEMPLATE,
    TABLE_HEADERS,
    TABLE_TITLE_TEMPLATE,
    TITLE,
    TOTAL_LABEL
)


class ReportGenerator:
    @staticmethod
    def generate_road_report(csv_path):
        """Генерация DOCX отчета из CSV данных"""
        df = pd.read_csv(csv_path)
        federal_roads = df[df['CLASSID'] == CLASSID_FEDERAL]
        regional_roads = df[df['CLASSID'] == CLASSID_REGIONAL]

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        title = doc.add_paragraph(TITLE)
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(12)
        intro = doc.add_paragraph(INTRO_TEXT)

        if federal_roads.empty:
            no_federal = doc.add_paragraph(NO_FEDERAL_TEXT)
            no_federal.runs[0].font.size = Pt(12)
        else:
            length_km = federal_roads['Shape_Length'].sum() / 1000
            with_federal_text = FEDERAL_PRESENT_TEXT.format(length=length_km)
            with_federal = doc.add_paragraph(with_federal_text)
            with_federal.runs[0].font.size = Pt(12)

        if not federal_roads.empty:
            ReportGenerator._add_roads_table(
                doc, federal_roads, "федерального значения"
            )

        if not regional_roads.empty:
            ReportGenerator._add_roads_table(
                doc, regional_roads,
                "регионального и межмуниципального значения"
            )

        return doc

    @staticmethod
    def _add_roads_table(doc, roads_df, road_type):
        """Добавляет таблицу с дорогами указанного типа"""
        table_title_text = TABLE_TITLE_TEMPLATE.format(road_type=road_type)
        table_title = doc.add_paragraph(f'\n{table_title_text}')
        table_title.runs[0].bold = True
        table_title.runs[0].font.size = Pt(12)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        col_widths = [Inches(0.35), Inches(1.85), Inches(2.6), Inches(1.14)]
        hdr_cells = table.rows[0].cells

        for i, header in enumerate(TABLE_HEADERS):
            cell = hdr_cells[i]
            cell.width = col_widths[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'

        total_length = 0
        for idx, row in roads_df.iterrows():
            row_cells = table.add_row().cells
            values = [
                str(idx + 1),
                str(row['ROAD_ID']) if pd.notna(row['ROAD_ID']) else '-',
                str(row['NAME']),
                f"{row['Shape_Length'] / 1000:.3f}"
            ]

            for i, value in enumerate(values):
                cell = row_cells[i]
                cell.width = col_widths[i]
                cell.text = value

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name ='Times New Roman'

            total_length += row['Shape_Length'] / 1000

        total_cells=table.add_row().cells
        total_cells[0].merge(total_cells[2])
        total_cells[0].text=TOTAL_LABEL
        total_cells[0].paragraphs[0].runs[0].bold=True
        total_cells[3].text=f"{total_length:.3f}"
        total_cells[3].paragraphs[0].runs[0].bold=True
        summary_text=SUMMARY_TEMPLATE.format(
            road_type=road_type,length=total_length
        )

        summary=doc.add_paragraph(summary_text)
        summary.runs[0].font.size=Pt(12)
